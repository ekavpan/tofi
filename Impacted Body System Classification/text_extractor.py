import pdfplumber
import pytesseract
import cv2
import numpy as np
from PIL import Image
import os
from typing import Optional
from docx import Document
from striprtf.striprtf import rtf_to_text
import platform
import subprocess
import pdf2image

# Try to import PyMuPDF, but don't fail if it's not available
try:
    import fitz
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    print("PyMuPDF not available. PDF extraction will use pdfplumber only.")

class TextExtractor:
    def __init__(self):
        """Initialize TextExtractor with Tesseract configuration"""
        # Set default Tesseract path for Windows
        if platform.system() == "Windows":
            default_paths = [
                r"C:\Program Files\Tesseract-OCR\tesseract.exe",
                r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
            ]
            for path in default_paths:
                if os.path.exists(path):
                    pytesseract.pytesseract.tesseract_cmd = path
                    print(f"Found Tesseract at: {path}")
                    break

    def test_tesseract(self) -> bool:
        """Test if Tesseract is working properly"""
        try:
            # Create a simple test image with text
            from PIL import Image, ImageDraw, ImageFont
            import numpy as np

            # Create a white image
            img = Image.new('RGB', (200, 50), color='white')
            d = ImageDraw.Draw(img)
            
            # Add text to image
            d.text((10,10), "Test 123", fill='black')
            
            # Try to extract text
            text = pytesseract.image_to_string(img)
            print(f"Tesseract test result: {text.strip()}")
            
            return "Test" in text or "123" in text
        except Exception as e:
            print(f"Tesseract test failed: {str(e)}")
            return False

    def _check_tesseract(self) -> bool:
        """Check if Tesseract is installed and accessible"""
        try:
            if platform.system() == "Windows":
                # Check if tesseract is in PATH or common install locations
                result = subprocess.run(['where', 'tesseract'], capture_output=True, text=True)
                if result.returncode == 0:
                    return True
                
                # Check common installation paths
                default_paths = [
                    r"C:\Program Files\Tesseract-OCR\tesseract.exe",
                    r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
                ]
                return any(os.path.exists(path) for path in default_paths)
            else:
                # For non-Windows systems
                result = subprocess.run(['which', 'tesseract'], capture_output=True)
                return result.returncode == 0
        except Exception:
            return False

    def _get_install_instructions(self) -> str:
        """Get platform-specific installation instructions"""
        if platform.system() == "Windows":
            return """
To install Tesseract OCR on Windows:
1. Download the installer from: https://github.com/UB-Mannheim/tesseract/wiki
2. Run the installer (select "Add to PATH" during installation)
3. Restart your application

Alternatively, you can install it via winget:
> winget install UB-Mannheim.TesseractOCR
"""
        else:
            return """
To install Tesseract OCR:
- Ubuntu/Debian: sudo apt-get install tesseract-ocr
- MacOS: brew install tesseract
- Other Linux: Check your package manager for tesseract-ocr
"""

    def preprocess_image(self, image):
        """Preprocess image to improve OCR accuracy"""
        # Convert to grayscale
        if len(image.shape) == 3:
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        else:
            gray = image

        # Apply thresholding to preprocess the image
        gray = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY | cv2.THRESH_OTSU)[1]

        # Apply dilation to connect text components
        kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (3,3))
        gray = cv2.dilate(gray, kernel, iterations=1)

        # Apply median blur to remove noise
        gray = cv2.medianBlur(gray, 3)

        return gray

    def extract_from_pdf(self, pdf_path: str) -> str:
        """Extract text from a PDF file with enhanced structure preservation"""
        try:
            print(f"\nAttempting to extract text from PDF: {pdf_path}")
            all_text = []
            
            # Method 1: Try pdfplumber first for better table extraction
            try:
                print("Attempting extraction with pdfplumber...")
                with pdfplumber.open(pdf_path) as pdf:
                    print(f"PDF has {len(pdf.pages)} pages")
                    
                    for page_num, page in enumerate(pdf.pages, 1):
                        print(f"\nProcessing page {page_num}")
                        page_text = []
                        
                        # First try to get raw text with layout preservation
                        text = page.extract_text(layout=True)
                        if text:
                            # Process text line by line
                            for line in text.split('\n'):
                                line = line.strip()
                                if line:
                                    # Keep lines with test results or important headers
                                    if any(c.isdigit() for c in line) or \
                                       any(keyword in line.upper() for keyword in ['TEST', 'RESULT', 'NAME', 'VALUE', 'UNIT', 'NORMAL', 'REFERENCE', 'RANGE']):
                                        page_text.append(line)
                                    # Keep other informative lines
                                    elif len(line) > 3 and not all(c in '.-_|\\/' for c in line):
                                        page_text.append(line)
                        
                        # Then try to extract tables
                        tables = page.extract_tables()
                        if tables:
                            print(f"Found {len(tables)} tables on page {page_num}")
                            for table in tables:
                                # Process each row in the table
                                for row in table:
                                    # Clean and join cells, preserving structure
                                    row_text = []
                                    for cell in row:
                                        if cell is not None:
                                            cell_text = str(cell).strip()
                                            if cell_text:
                                                # Always keep cells with numbers or units
                                                if any(c.isdigit() for c in cell_text) or \
                                                   any(unit in cell_text.upper() for unit in ['MG', 'G', 'ML', 'L', 'MMOL', 'UNITS']):
                                                    row_text.append(cell_text)
                                                # Keep other non-empty cells
                                                elif len(cell_text) > 1:
                                                    row_text.append(cell_text)
                                    if row_text:
                                        line = " | ".join(row_text)
                                        if line not in page_text:  # Avoid duplicates
                                            page_text.append(line)
                        
                        # Add page separator and page text
                        if page_text:
                            all_text.extend([f"\n=== Page {page_num} ===\n"])
                            all_text.extend(page_text)
                            print(f"Extracted {len(page_text)} lines from page {page_num}")
                    
                    if all_text:
                        # Post-process extracted text
                        seen = set()
                        final_lines = []
                        current_section = None
                        
                        for line in all_text:
                            line = line.strip()
                            if not line or (line in seen and not line.startswith("===")):
                                continue
                            
                            if not line.startswith("==="):
                                seen.add(line)
                            
                            # Add section headers
                            if line.isupper() and len(line) > 3 and not line.startswith("==="):
                                if current_section != line:
                                    current_section = line
                                    final_lines.extend(['', line, '-' * len(line)])
                            # Add test results and other content
                            else:
                                final_lines.append(line)
                        
                        final_text = '\n'.join(final_lines)
                        print("\nExtracted text preview:")
                        print("First 200 chars:", final_text[:200])
                        print("Last 200 chars:", final_text[-200:])
                        return final_text
                
            except Exception as e:
                print(f"pdfplumber extraction failed: {str(e)}")
                raise
            
        except Exception as e:
            print(f"Error extracting text from PDF: {str(e)}")
            raise

    def extract_from_image(self, image_path: str) -> str:
        """Extract text from an image using OCR"""
        try:
            if not self._check_tesseract():
                raise Exception(
                    f"Tesseract is not installed or not in PATH.\n"
                    f"{self._get_install_instructions()}"
                )
            
            # Read image using OpenCV
            image = cv2.imread(image_path)
            if image is None:
                raise Exception("Failed to load image")

            # Preprocess the image
            processed_image = self.preprocess_image(image)

            # Extract text using Tesseract
            text = pytesseract.image_to_string(
                Image.fromarray(processed_image),
                config='--psm 6'  # Assume uniform block of text
            )

            if not text.strip():
                raise Exception(
                    "No text was extracted from the image. This could be because:\n"
                    "1. The image doesn't contain text\n"
                    "2. The text is not clear enough\n"
                    "3. The image format is not supported\n"
                    "Please try uploading a clearer image or a different format."
                )

            return text

        except Exception as e:
            raise Exception(f"Error extracting text from image: {str(e)}")

    def extract_text(self, file_path: str) -> str:
        """Extract text from various file types with improved handling"""
        try:
            # Get file extension
            _, ext = os.path.splitext(file_path)
            ext = ext.lower()
            
            if ext == '.pdf':
                return self.extract_from_pdf(file_path)
            elif ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read().strip()
            elif ext == '.docx':
                doc = Document(file_path)
                return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            elif ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
                return self.extract_from_image(file_path)
            else:
                raise ValueError(f"Unsupported file format: {ext}")
                
        except UnicodeDecodeError:
            # Try alternative encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    return f.read().strip()
            except Exception as e:
                raise ValueError(f"Error reading file with alternative encoding: {str(e)}")
        except Exception as e:
            raise ValueError(f"Error extracting text from file: {str(e)}")

    def extract_from_text_file(self, file_path: str) -> str:
        """Extract text from text-based files"""
        try:
            # Handle .txt files
            if file_path.lower().endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read().strip()
            
            # Handle .doc, .docx files
            elif file_path.lower().endswith(('.doc', '.docx')):
                doc = Document(file_path)
                return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            
            # Handle .rtf files
            elif file_path.lower().endswith('.rtf'):
                with open(file_path, 'r', encoding='utf-8') as f:
                    rtf = f.read()
                    return rtf_to_text(rtf).strip()
                    
        except Exception as e:
            raise Exception(f"Error extracting text from file: {str(e)}")
